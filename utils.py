# utils.py
import os
import re
from typing import Dict, Any, List, Tuple
from docx import Document
from io import BytesIO
from agents import ResearcherAgent, WriterAgent, CriticAgent
import streamlit as st
import requests

def generate_outline(target: str, researcher: ResearcherAgent, writer: WriterAgent, critic: CriticAgent, max_iterations: int = 4, iteration: int = 0) -> Dict:
    if iteration >= max_iterations:
        return {}
    # Use agents to generate and refine outline
    research = researcher.research(target)
    outline = writer.write(research)
    feedback = critic.critique(outline)
    # If feedback is satisfactory, return outline
    if "no further improvements" in feedback.lower():
        return {"Title": target, "Outline": outline}
    else:
        # Recursive call with updated feedback
        return generate_outline(target, researcher, writer, critic, max_iterations, iteration + 1)

def create_document(outline: Dict, researcher: ResearcherAgent, writer: WriterAgent, critic: CriticAgent, max_iterations: int = 4) -> Document:
    doc = Document()
    doc.add_heading(outline.get('Title', 'Research Paper'), 0)
    sections = outline.get('Outline', '').split('\n')
    for section in sections:
        process_paragraph(doc, section, researcher, writer, critic, max_iterations)
    return doc

def process_paragraph(doc: Document, title: str, researcher: ResearcherAgent, writer: WriterAgent, critic: CriticAgent, max_iterations: int = 4, iteration: int = 0):
    if iteration >= max_iterations:
        return
    research = researcher.research(title)
    content = writer.write(research)
    feedback = critic.critique(content)
    if "no further improvements" in feedback.lower():
        doc.add_heading(title, level=1)
        doc.add_paragraph(content)
    else:
        # Recursive call with updated feedback
        process_paragraph(doc, title, researcher, writer, critic, max_iterations, iteration + 1)

def cleanup_and_format_document(doc: Document) -> Document:
    # Apply consistent formatting
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = docx.shared.Pt(12)
    # Add page numbers
    section = doc.sections[0]
    footer = section.footer
    footer.paragraphs[0].text = "Page " + docx.oxml.shared.OxmlElement('w:pageNum')
    # Insert Table of Contents
    doc.add_page_break()
    doc.add_heading('Table of Contents', level=1)
    # Note: Implement TOC if necessary
    return doc

def format_tables(doc: Document) -> Document:
    for table in doc.tables:
        table.style = 'Table Grid'
        # Format header row
        for cell in table.rows[0].cells:
            for paragraph in cell.paragraphs:
                paragraph.runs[0].bold = True
        # Adjust column widths, add captions
    return doc

def process_figures(doc: Document) -> Document:
    # Find and replace figure placeholders
    for para in doc.paragraphs:
        if '$$' in para.text:
            description = para.text.strip('$$')
            image = generate_figure(description)
            if image:
                para.text = ''
                para.add_run().add_picture(image)
                para.add_run(f"Figure: {description}").italic = True
    return doc

def generate_figure(description: str):
    # Use Napkin.ai API to generate figure
    napkin_api_key = os.environ.get('NAPKIN_AI_API_KEY')
    # Implement API call to generate figure
    # Return image path or BytesIO object
    return None  # Placeholder

def manage_references(doc: Document, citation_styles: List[str]) -> Document:
    # Collect references, replace in-text citations
    references = []
    # Implement reference management
    doc.add_heading('References', level=1)
    for ref in references:
        doc.add_paragraph(ref)
    return doc

def finalize_for_ms_word(doc: Document) -> Document:
    # Ensure compatibility
    return doc

def assemble_final_document(doc: Document) -> Document:
    # Assemble sections in correct order
    return doc

def perform_quality_check(doc: Document) -> Tuple[Document, List[str]]:
    warnings = []
    # Perform checks
    return doc, warnings

def save_and_prepare_download(doc: Document) -> BytesIO:
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream
